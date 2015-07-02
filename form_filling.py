import docx
from docx import table

def form_filling(date, name, amt,reg_no):
	doc = docx.Document('E-receiptTemplate2015.docx')
	table = doc.tables[0]

	table.rows[0].cells[1].paragraphs[0].text = reg_no
	table.rows[1].cells[1].paragraphs[0].text = name
	table.rows[2].cells[1].paragraphs[0].text = "Delegate"
	table.rows[3].cells[1].paragraphs[0].text = amt
	table.rows[4].cells[1].paragraphs[0].text = "online"
	table.rows[5].cells[1].paragraphs[0].text = date
	filename = doc.save(reg_no +" " +name+".docx")











	